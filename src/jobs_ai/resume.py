from pathlib import Path
from typing import Iterable, Optional

from jobs_ai.models import ResumeDocument
from jobs_ai.preferences import detect_language


class ResumeStore:
    supported_suffixes = {".md", ".txt", ".pdf", ".docx"}

    def __init__(self, path: Path, cv_directory: Path) -> None:
        self.path = path
        self.cv_directory = cv_directory
        self.path.parent.mkdir(parents=True, exist_ok=True)

    def read(self) -> str:
        documents = self.read_documents()
        if documents:
            return "\n\n".join(_format_document(document) for document in documents)
        if not self.path.exists():
            return ""
        return self.path.read_text(encoding="utf-8").strip()

    def read_documents(self) -> list[ResumeDocument]:
        documents = []
        for file_path in self._candidate_files():
            content = _read_resume_file(file_path)
            if content:
                documents.append(
                    ResumeDocument(
                        name=file_path.name,
                        path=str(file_path),
                        language=detect_language(content),
                        content=content,
                    )
                )
        if documents:
            return documents
        if self.path.exists():
            content = self.path.read_text(encoding="utf-8").strip()
            if content:
                return [
                    ResumeDocument(
                        name=self.path.name,
                        path=str(self.path),
                        language=detect_language(content),
                        content=content,
                    )
                ]
        return []

    def select_for_language(self, language: str) -> Optional[ResumeDocument]:
        documents = self.read_documents()
        if not documents:
            return None
        if language != "unknown":
            for document in documents:
                if document.language == language:
                    return document
            for document in documents:
                if language in document.name.lower():
                    return document
        return documents[0]

    def write(self, content: str) -> None:
        self.path.write_text(content.strip() + "\n", encoding="utf-8")

    def _candidate_files(self) -> Iterable[Path]:
        if not self.cv_directory.exists():
            return []
        return sorted(
            path
            for path in self.cv_directory.rglob("*")
            if path.is_file() and path.suffix.lower() in self.supported_suffixes
        )


def _read_resume_file(path: Path) -> str:
    suffix = path.suffix.lower()
    try:
        if suffix in {".md", ".txt"}:
            return path.read_text(encoding="utf-8", errors="ignore").strip()
        if suffix == ".pdf":
            from pypdf import PdfReader

            reader = PdfReader(str(path))
            return "\n".join(page.extract_text() or "" for page in reader.pages).strip()
        if suffix == ".docx":
            from docx import Document

            document = Document(str(path))
            return "\n".join(paragraph.text for paragraph in document.paragraphs).strip()
    except Exception:
        return ""
    return ""


def _format_document(document: ResumeDocument) -> str:
    return (
        f"## CV: {document.name}\n"
        f"Path: {document.path}\n"
        f"Detected language: {document.language}\n\n"
        f"{document.content}"
    )
